import os
import shutil
import sys
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from tkinter import BOTH, DISABLED, NORMAL, Button, Frame, Label, PhotoImage, Tk, Toplevel, messagebox

import mss
import mss.tools
from docx import Document
from docx.shared import Inches


APP_NAME = "Screenshot Saver"


@dataclass
class ScreenshotRecord:
    number: int
    captured_at: datetime
    path: Path


def get_app_folder() -> Path:
    """Return the folder beside the .exe, or beside this script during development."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def get_resource_path(relative_path: str) -> Path:
    """Return a bundled PyInstaller resource path, or a development file path."""
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS) / relative_path
    return Path(__file__).resolve().parent / relative_path


class ScreenshotSaverApp:
    def __init__(self, root: Tk) -> None:
        self.root = root
        self.root.title(APP_NAME)
        self.root.geometry("420x320")
        self.root.resizable(False, False)

        self.app_folder = get_app_folder()
        self.temp_folder = self.app_folder / "screenshot_temp"
        self.records: list[ScreenshotRecord] = []
        self.session_started = False
        self.background_image: PhotoImage | None = None
        self.camera_image: PhotoImage | None = None
        self.pause_image: PhotoImage | None = None

        self._build_ui()
        self.root.protocol("WM_DELETE_WINDOW", self.exit_app)

    def _build_ui(self) -> None:
        self._load_button_images()

        if self.background_image:
            background_label = Label(self.root, image=self.background_image, borderwidth=0)
            background_label.place(x=0, y=0, relwidth=1, relheight=1)

        container = Frame(self.root, padx=22, pady=18, bg="#f7fbff")
        container.place(x=28, y=24, width=364, height=272)

        title_label = Label(
            container,
            text=APP_NAME,
            bg="#f7fbff",
            fg="#183153",
            font=("Segoe UI", 16, "bold"),
        )
        title_label.pack(pady=(0, 8))

        self.status_label = Label(
            container,
            text="Ready. Click Camera to take a screenshot.",
            anchor="center",
            justify="center",
            wraplength=310,
            width=42,
            height=3,
            bg="#edf6ff",
            fg="#183153",
            font=("Segoe UI", 9),
            relief="flat",
        )
        self.status_label.pack(pady=(0, 12), fill="x")

        button_row = Frame(container, bg="#f7fbff")
        button_row.pack(pady=2)

        if self.camera_image and self.pause_image:
            self.camera_button = Button(
                button_row,
                image=self.camera_image,
                width=112,
                height=90,
                borderwidth=0,
                cursor="hand2",
                bg="#f7fbff",
                activebackground="#e7f1ff",
                command=self.take_screenshot,
            )
            self.pause_button = Button(
                button_row,
                image=self.pause_image,
                width=112,
                height=90,
                borderwidth=0,
                cursor="hand2",
                bg="#f7fbff",
                activebackground="#e7f1ff",
                state=DISABLED,
                command=self.pause_stop,
            )
        else:
            self.camera_button = Button(
                button_row,
                text="Camera",
                width=14,
                command=self.take_screenshot,
            )
            self.pause_button = Button(
                button_row,
                text="Pause / Stop",
                width=14,
                state=DISABLED,
                command=self.pause_stop,
            )

        self.camera_button.pack(side="left", padx=5)
        self.pause_button.pack(side="left", padx=5)

        self.exit_button = Button(
            container,
            text="Exit",
            width=18,
            bg="#183153",
            fg="white",
            activebackground="#28466f",
            activeforeground="white",
            cursor="hand2",
            command=self.exit_app,
        )
        self.exit_button.pack(side="bottom", pady=(12, 0))

    def _load_button_images(self) -> None:
        try:
            self.background_image = PhotoImage(file=get_resource_path("assets/background_app.png"))
        except Exception:
            self.background_image = None

        try:
            self.camera_image = PhotoImage(file=get_resource_path("assets/camera.png")).subsample(3, 3)
            self.pause_image = PhotoImage(file=get_resource_path("assets/pause.png")).subsample(3, 3)
        except Exception:
            self.camera_image = None
            self.pause_image = None

    def set_status(self, text: str) -> None:
        self.root.after(0, lambda: self.status_label.config(text=text))

    def take_screenshot(self) -> None:
        if not self.session_started:
            self._prepare_temp_folder()
            self.records = []
            self.session_started = True
            self.pause_button.config(state=NORMAL)

        self.camera_button.config(state=DISABLED)
        self.set_status("Taking screenshot...")

        try:
            self._capture_once()
            self.set_status(f"Screenshot {len(self.records)} saved. Move to the next page and click Camera again.")
        except Exception as exc:
            self.set_status(f"Screenshot failed: {exc}")
        finally:
            self.camera_button.config(state=NORMAL)

    def _prepare_temp_folder(self) -> None:
        self.temp_folder.mkdir(parents=True, exist_ok=True)
        for item in self.temp_folder.iterdir():
            if item.is_file():
                item.unlink()
            elif item.is_dir():
                shutil.rmtree(item)

    def _capture_once(self) -> None:
        screenshot_number = len(self.records) + 1
        captured_at = datetime.now()
        filename = f"screenshot_{screenshot_number:03d}_{captured_at:%Y-%m-%d_%H-%M-%S}.png"
        output_path = self.temp_folder / filename

        self.root.withdraw()
        self.root.update()
        time.sleep(0.25)

        try:
            with mss.mss() as screen_capture:
                monitor = screen_capture.monitors[0]
                image = screen_capture.grab(monitor)
                mss.tools.to_png(image.rgb, image.size, output=str(output_path))
        finally:
            self.root.deiconify()
            self.root.lift()
            self.root.focus_force()

        self.records.append(
            ScreenshotRecord(
                number=screenshot_number,
                captured_at=captured_at,
                path=output_path,
            )
        )

    def pause_stop(self) -> None:
        if not self.records:
            self.set_status("Paused. No screenshots were captured.")
            self.session_started = False
            self.pause_button.config(state=DISABLED)
            return

        self.set_status("Creating Word file...")
        self.camera_button.config(state=DISABLED)
        self.pause_button.config(state=DISABLED)

        try:
            docx_path = self._create_word_document()
            self._cleanup_temp_folder()
            self.records = []
            self.session_started = False
            self.set_status(f"Paused. Word file saved:\n{docx_path.name}")
            self.show_saved_popup(docx_path)
            self.open_saved_document(docx_path)
        except Exception as exc:
            self.set_status(f"Could not create Word file: {exc}")
        finally:
            self.camera_button.config(state=NORMAL)

    def open_saved_document(self, docx_path: Path) -> None:
        os.startfile(docx_path)

    def show_saved_popup(self, docx_path: Path) -> None:
        popup = Toplevel(self.root)
        popup.title(APP_NAME)
        popup.resizable(False, False)
        popup.attributes("-topmost", True)

        width = 330
        height = 110
        screen_width = popup.winfo_screenwidth()
        screen_height = popup.winfo_screenheight()
        x = screen_width - width - 18
        y = screen_height - height - 58
        popup.geometry(f"{width}x{height}+{x}+{y}")

        frame = Frame(popup, padx=14, pady=12)
        frame.pack(fill=BOTH, expand=True)

        Label(frame, text="Word file saved", font=("Segoe UI", 11, "bold")).pack(anchor="w")
        Label(frame, text=docx_path.name, wraplength=292, justify="left").pack(anchor="w", pady=(6, 10))
        Button(frame, text="OK", width=10, command=popup.destroy).pack(anchor="e")

        popup.after(6000, popup.destroy)

    def _create_word_document(self) -> Path:
        document = Document()
        document.add_heading("Screenshots", level=1)

        for record in self.records:
            document.add_heading(f"Screenshot {record.number}", level=2)
            document.add_paragraph(f"Date and time: {record.captured_at:%Y-%m-%d %H:%M:%S}")
            document.add_picture(str(record.path), width=Inches(6.5))

        output_name = f"Screenshots_{datetime.now():%Y-%m-%d_%H-%M-%S}.docx"
        output_path = self.app_folder / output_name
        document.save(output_path)
        return output_path

    def _cleanup_temp_folder(self) -> None:
        if self.temp_folder.exists():
            shutil.rmtree(self.temp_folder)

    def exit_app(self) -> None:
        if self.records:
            should_exit = messagebox.askyesno(
                APP_NAME,
                "You have screenshots that are not in a Word file yet. Create the Word file before exiting?",
            )
            if should_exit:
                self.pause_stop()
                return

        self.root.destroy()


def main() -> None:
    root = Tk()
    ScreenshotSaverApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
